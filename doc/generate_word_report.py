# -*- coding: utf-8 -*-
"""Generate Word report with charts from PROJECT_STATUS_AND_ISSUES.md content."""

import os
import tempfile
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml import OxmlElement

# Chinese font for matplotlib
plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "项目现状问题分析与解决方案.docx")
CHART_DIR = tempfile.mkdtemp()


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], "2E75B6")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(10)
        if ri % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "DEEAF6")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph()
    return table


def save_chart(fig, name):
    path = os.path.join(CHART_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def add_chart(doc, path, caption, width=6.0):
    doc.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()


def chart_data_pipeline():
    stages = ["原始合并", "去重后", "分层抽样", "训练集", "测试集"]
    counts = [183239, 167308, 5976, 4780, 1196]
    colors = ["#4472C4", "#5B9BD5", "#70AD47", "#FFC000", "#ED7D31"]
    fig, ax = plt.subplots(figsize=(9, 5))
    bars = ax.bar(stages, counts, color=colors, edgecolor="white", linewidth=0.8)
    ax.set_ylabel("记录数", fontsize=11)
    ax.set_title("数据处理流水线规模", fontsize=13, fontweight="bold", pad=12)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{int(x):,}"))
    for bar, val in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 3000,
                f"{val:,}", ha="center", va="bottom", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return save_chart(fig, "data_pipeline.png")


def chart_model_accuracy():
    snr_levels = ["Clean", "5dB", "0dB", "-5dB"]
    lgbm = [0.48, 0.08, 0.07, 0.07]
    yamnet = [1.91, 0.42, 0.37, 0.12]
    baseline = 0.122
    x = range(len(snr_levels))
    w = 0.3
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.bar([i - w for i in x], lgbm, w, label="LightGBM", color="#4472C4")
    ax.bar([i for i in x], yamnet, w, label="YAMNet（冻结）", color="#70AD47")
    ax.axhline(y=baseline, color="#C00000", linestyle="--", linewidth=1.5, label=f"随机基线 ({baseline}%)")
    ax.set_xticks(list(x))
    ax.set_xticklabels(snr_levels)
    ax.set_ylabel("准确率 (%)", fontsize=11)
    ax.set_title("不同信噪比条件下模型准确率对比", fontsize=13, fontweight="bold", pad=12)
    ax.legend(loc="upper right", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return save_chart(fig, "model_accuracy.png")


def chart_fold_results():
    folds = [1, 2, 3, 4, 5]
    lgbm_clean = [0.59, 0.67, 0.42, 0.50, 0.25]
    yamnet_clean = [1.51, 2.09, 1.51, 2.17, 2.26]
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.plot(folds, lgbm_clean, "o-", color="#4472C4", linewidth=2, markersize=8, label="LightGBM (Clean)")
    ax.plot(folds, yamnet_clean, "s-", color="#70AD47", linewidth=2, markersize=8, label="YAMNet (Clean)")
    ax.set_xlabel("折数 (Fold)", fontsize=11)
    ax.set_ylabel("Clean 准确率 (%)", fontsize=11)
    ax.set_title("5 折交叉验证 Clean 准确率逐折对比", fontsize=13, fontweight="bold", pad=12)
    ax.set_xticks(folds)
    ax.legend(fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return save_chart(fig, "fold_clean.png")


def chart_snr_decay():
    snr_levels = ["Clean", "5dB", "0dB", "-5dB"]
    lgbm = [0.48, 0.08, 0.07, 0.07]
    yamnet = [1.91, 0.42, 0.37, 0.12]
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.plot(snr_levels, lgbm, "o-", color="#4472C4", linewidth=2.5, markersize=9, label="LightGBM")
    ax.plot(snr_levels, yamnet, "s-", color="#70AD47", linewidth=2.5, markersize=9, label="YAMNet（冻结）")
    ax.axhline(y=0.122, color="#C00000", linestyle="--", linewidth=1.5, label="随机基线 (0.122%)")
    ax.set_ylabel("准确率 (%)", fontsize=11)
    ax.set_title("噪声衰减曲线：准确率随 SNR 下降的变化", fontsize=13, fontweight="bold", pad=12)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return save_chart(fig, "snr_decay.png")


def chart_inference_latency():
    models = ["LightGBM", "YAMNet"]
    latency = [154.5, 86.0]
    colors = ["#4472C4", "#70AD47"]
    fig, ax = plt.subplots(figsize=(7, 4.5))
    bars = ax.barh(models, latency, color=colors, height=0.5)
    ax.set_xlabel("推理延迟 (ms)", fontsize=11)
    ax.set_title("推理延迟对比", fontsize=13, fontweight="bold", pad=12)
    for bar, val in zip(bars, latency):
        ax.text(val + 3, bar.get_y() + bar.get_height() / 2, f"{val} ms",
                va="center", fontsize=10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return save_chart(fig, "inference.png")


def chart_problem_overview():
    layers = ["数据", "模型架构", "训练策略", "评估方法", "LightGBM", "代码 Bug"]
    counts = [5, 4, 4, 3, 4, 8]
    colors = ["#4472C4", "#5B9BD5", "#70AD47", "#FFC000", "#ED7D31", "#A5A5A5"]
    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.barh(layers, counts, color=colors, height=0.6)
    ax.set_xlabel("问题数量", fontsize=11)
    ax.set_title("各层面发现问题数量", fontsize=13, fontweight="bold", pad=12)
    for bar, val in zip(bars, counts):
        ax.text(val + 0.15, bar.get_y() + bar.get_height() / 2, str(val),
                va="center", fontsize=10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return save_chart(fig, "problems.png")


def chart_code_bug_status():
    labels = ["已修复", "待处理"]
    sizes = [8, 0]
    colors = ["#70AD47", "#C00000"]
    explode = (0.05, 0)
    fig, ax = plt.subplots(figsize=(6, 5))
    wedges, texts, autotexts = ax.pie(
        sizes, explode=explode, labels=labels, colors=colors,
        autopct="%1.0f%%", startangle=90, textprops={"fontsize": 11}
    )
    ax.set_title("YAMNet 端到端代码问题修复状态（共 8 项）", fontsize=12, fontweight="bold", pad=12)
    plt.tight_layout()
    return save_chart(fig, "bug_status.png")


def chart_yamnet_advantage():
    metrics = ["Clean 准确率", "5dB 准确率", "0dB 准确率", "-5dB 准确率", "Clean F1"]
    advantage = [4.0, 5.0, 5.5, 1.7, 4.0]
    fig, ax = plt.subplots(figsize=(9, 5))
    bars = ax.bar(metrics, advantage, color="#70AD47", edgecolor="white")
    ax.set_ylabel("YAMNet 相对 LightGBM 倍数", fontsize=11)
    ax.set_title("YAMNet 相对 LightGBM 的性能优势", fontsize=13, fontweight="bold", pad=12)
    ax.axhline(y=1, color="#888888", linestyle="--", linewidth=1)
    for bar, val in zip(bars, advantage):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                f"{val}x", ha="center", va="bottom", fontsize=9)
    plt.xticks(rotation=15, ha="right")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    return save_chart(fig, "yamnet_advantage.png")


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("项目现状、问题分析与解决方案")
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("噪声鲁棒的鸟类发声分类——面向低资源场景")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("更新日期：2026年7月16日")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()
    add_para(doc, "本文档汇总当前项目进展、发现的各类问题、建议的修复方案，以及需要与教授讨论确认的开放问题。")

    # Section 1
    add_heading(doc, "一、项目现状概述", 1)

    add_heading(doc, "1.1 项目基本信息", 2)
    add_table(doc,
              ["项目", "说明"],
              [
                  ["名称", "噪声鲁棒的鸟类发声分类——面向低资源场景"],
                  ["核心任务", "从野外录音中识别 1,229 个鸟类物种，面临长尾分布和复杂环境噪声"],
                  ["团队", "3 人（张佳男 / 陈锦城 / 黄文娟）"],
                  ["数据来源", "BirdCLEF 2021-2026 竞赛合集（Xeno-Canto + iNaturalist）"],
              ])

    add_heading(doc, "1.2 数据规模", 2)
    add_table(doc,
              ["阶段", "记录数", "说明"],
              [
                  ["原始合并", "183,239", "6 个年度 CSV 合并"],
                  ["字段级去重后", "167,308", "零信息损失智能合并"],
                  ["分层抽样子集", "5,976", "双层分层（物种 × SNR），每物种 ≥ 5 样本"],
                  ["训练集", "4,780（80%）", "按 SNR 分层划分"],
                  ["留出测试集", "1,196（20%）", "按 SNR 分层划分"],
                  ["交叉验证", "5 折", "训练 3,824 / 验证 956"],
              ])
    add_chart(doc, chart_data_pipeline(), "图 1  数据处理流水线各阶段记录数")

    add_heading(doc, "1.3 三种模型方案", 2)
    add_table(doc,
              ["模型", "负责人", "方法", "状态"],
              [
                  ["LightGBM", "张佳男", "手工表格特征 + 梯度提升树", "5 折 CV + 噪声评估已完成"],
                  ["FastAI CNN", "陈锦城", "梅尔频谱图图像分类", "待完成"],
                  ["YAMNet 冻结", "黄文娟", "预训练嵌入 + 分类头", "5 折 CV + 噪声评估已完成"],
                  ["YAMNet 端到端微调", "黄文娟", "解冻顶层 + MixUp + 差分学习率 + 类别权重", "代码已完成并修复，待 Kaggle 实跑"],
              ])

    # Section 2
    add_heading(doc, "二、当前实验结果", 1)

    add_heading(doc, "2.1 已完成模型结果对比", 2)
    add_table(doc,
              ["指标", "LightGBM", "YAMNet（冻结）", "YAMNet 优势"],
              [
                  ["Clean 准确率", "0.48% ± 0.14%", "1.91% ± 0.33%", "4.0 倍"],
                  ["5dB 准确率", "0.08% ± 0.05%", "0.42% ± 0.19%", "5.0 倍"],
                  ["0dB 准确率", "0.07% ± 0.03%", "0.37% ± 0.15%", "5.5 倍"],
                  ["-5dB 准确率", "0.07% ± 0.03%", "0.12% ± 0.04%", "1.7 倍"],
                  ["Clean F1 (weighted)", "0.0045", "0.0178", "4.0 倍"],
                  ["推理延迟", "154.5 ms", "86.0 ms", "YAMNet 快 44%"],
                  ["GPU 显存峰值", "N/A（CPU）", "211.4 MB", "—"],
              ])
    add_para(doc, "随机基线 = 1/818 = 0.122%（测试集 818 个类）。")
    add_chart(doc, chart_model_accuracy(), "图 2  不同信噪比条件下模型准确率对比")
    add_chart(doc, chart_snr_decay(), "图 3  噪声衰减曲线")
    add_chart(doc, chart_yamnet_advantage(), "图 4  YAMNet 相对 LightGBM 的性能优势倍数")
    add_chart(doc, chart_inference_latency(), "图 5  推理延迟对比")

    add_heading(doc, "2.2 YAMNet（冻结）5 折逐折结果", 2)
    add_table(doc,
              ["Fold", "Clean", "5dB", "0dB", "-5dB"],
              [
                  ["1", "1.51%", "0.67%", "0.42%", "0.08%"],
                  ["2", "2.09%", "0.42%", "0.42%", "0.17%"],
                  ["3", "1.51%", "0.17%", "0.42%", "0.08%"],
                  ["4", "2.17%", "0.25%", "0.08%", "0.08%"],
                  ["5", "2.26%", "0.59%", "0.50%", "0.17%"],
              ])

    add_heading(doc, "2.3 LightGBM 5 折逐折结果", 2)
    add_table(doc,
              ["Fold", "Clean", "5dB", "0dB", "-5dB"],
              [
                  ["1", "0.59%", "0.08%", "0.00%", "0.08%"],
                  ["2", "0.67%", "0.17%", "0.08%", "0.00%"],
                  ["3", "0.42%", "0.00%", "0.08%", "0.08%"],
                  ["4", "0.50%", "0.08%", "0.08%", "0.08%"],
                  ["5", "0.25%", "0.08%", "0.08%", "0.08%"],
              ])
    add_chart(doc, chart_fold_results(), "图 6  5 折交叉验证 Clean 准确率逐折对比")

    add_heading(doc, "2.4 关键发现", 2)
    findings = [
        "YAMNet 在所有 SNR 档位均优于 LightGBM（4.0 倍至 5.5 倍），符合迁移学习抗噪能力强的预期。",
        "两者在 0dB 及以下基本塌缩为随机水平，与文献中 BirdNET「SNR < 3dB 时显著下降」的结论一致。",
        "LightGBM 在 5dB 就已低于随机基线（0.08% < 0.122%），说明手工特征对噪声极其敏感。",
        "绝对准确率低是数据约束决定的：1,229 类 × 每类仅约 3.9 个训练样本，测试集 63% 的类只有 1 个样本。但两模型均显著优于随机（LightGBM 4 倍随机，YAMNet 15.6 倍随机），证明模型学到了有效特征。",
        "YAMNet 推理反而更快（86ms vs 154ms），因为 LightGBM 的手工特征提取开销大于 YAMNet 嵌入提取。",
    ]
    for f in findings:
        add_bullet(doc, f)

    # Section 3
    add_heading(doc, "三、发现的问题", 1)

    add_heading(doc, "3.1 数据层面问题", 2)
    data_issues = [
        ("D1：极端的类别数与样本数比例",
         "1,229 类 × 训练集 4,780 条 = 平均每类仅约 3.9 个样本。5 折交叉验证后，每折训练集 3,824 条 / 1,229 类 = 每折平均仅约 3.1 个样本/类。部分稀有物种在某一折的训练集里可能只有 1-2 个样本。影响：模型几乎不可能学到这些类的判别特征，准确率被严重拖累。"),
        ("D2：测试集的「单样本类」问题",
         "测试集 1,196 条 / 1,229 类 = 平均每类不到 1 条。约 63% 的类在测试集中只有 1 个样本。预测对为 100%，错为 0%，造成评估方差极大，几个样本翻转就显著改变整体数值。"),
        ("D3：弱标签（secondary_labels）未利用",
         "录音中可能同时出现多种鸟，secondary_labels 记录了这些信息。当前训练和评估都只使用 primary_label，丢弃了多标签信息，导致标签噪声。"),
        ("D4：SNR 代理指标不够精确",
         "用 rating（0-5 主观质量评分）作为 SNR 代理。同样 rating=3 的两条录音，实际信噪比可能差异巨大，分层抽样的 SNR 分布可能不如预期均匀。"),
        ("D5：长尾分布未被有效平衡",
         "「每物种 ≥ 5 样本」保证全覆盖，但分布仍然长尾。高频物种 10+ 样本，低频物种恰好 5 个。冻结策略阶段训练时没有用 class weight 或 focal loss 补偿。"),
    ]
    for title, desc in data_issues:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)

    add_heading(doc, "3.2 模型架构层面问题", 2)
    arch_issues = [
        ("A1：YAMNet 域不匹配", "YAMNet 在 AudioSet（521 类通用音频）上预训练，鸟鸣声是 AudioSet 中的小子集，嵌入空间对鸟鸣的细粒度区分能力有限。"),
        ("A2：嵌入聚合方式过于简单", "当前对所有帧取平均值，鸟鸣的核心判别信息往往在特定时间段，平均池化把关键信息稀释了。"),
        ("A3：分类头容量不足（冻结策略）", "Dense(256, ReLU) → Dropout(0.3) → Dense(1229, Softmax)，单层全连接网络难以学习 1229 类的复杂决策边界。"),
        ("A4：冻结编码器无法适配鸟鸣", "YAMNet 全部参数冻结，只调分类头，特征提取层完全无法适应鸟鸣的频谱特征。"),
    ]
    for title, desc in arch_issues:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)

    add_heading(doc, "3.3 训练策略层面问题", 2)
    train_issues = [
        ("T1：冻结策略完全没有数据增强", "无 MixUp、无 SpecAugment、无时间/频率掩蔽、无波形增强。在每类仅 3 个样本的情况下，不做增强等于放弃了最有效的正则化手段。"),
        ("T2：冻结策略没有类别平衡损失", "使用标准 sparse_categorical_crossentropy，所有样本权重相同，长尾分布下头部类梯度主导更新方向。"),
        ("T3：固定 5 秒中段截取可能丢失关键信息", "鸟鸣可能在前 2 秒或后 2 秒，某些录音有效鸟鸣段很短，被稀释在 5 秒中段里。"),
        ("T4：交叉验证在极小类上的不稳定性", "不同 fold 结果差异大（fold1=1.51% vs fold5=2.26%），说明少量样本的随机分配显著影响学习效果。"),
    ]
    for title, desc in train_issues:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)

    add_heading(doc, "3.4 评估方法层面问题", 2)
    eval_issues = [
        ("E1：Top-1 准确率对 1229 类过于严苛", "细粒度分类文献通常用 Top-5 准确率或 macro-F1 作为主要指标，当前只报告了 Top-1 accuracy 和 weighted-F1。"),
        ("E2：没有按类别频率分组的细粒度分析", "只有总体准确率，没有拆分「高频物种 vs 低频物种」的性能，聚合指标掩盖了部分类别可能已经学到的事实。"),
        ("E3：缺少置信度分析", "缺少 Top-k 准确率来展示模型实际学到了多少。"),
    ]
    for title, desc in eval_issues:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)

    add_heading(doc, "3.5 LightGBM 特有问题", 2)
    add_para(doc, "L1：手工特征对细粒度分类几乎无效", bold=True)
    add_para(doc, "输入是频谱质心、带宽、MFCC 统计值、过零率等摘要统计量，1229 种鸟类的区分靠的是声音的精细时频结构，手工特征根本无法表达。", indent=True)
    add_para(doc, "L2：没有任何迁移学习或先验知识", bold=True)
    add_table(doc,
              ["模型", "先验知识"],
              [
                  ["YAMNet", "AudioSet 521 类预训练（含鸟鸣子类）"],
                  ["FastAI CNN", "ImageNet 预训练（视觉特征可迁移到频谱图）"],
                  ["LightGBM", "完全从零开始，无任何先验"],
              ])
    add_para(doc, "L3：手工特征对噪声极其敏感", bold=True)
    add_table(doc,
              ["特征", "噪声敏感性", "原因"],
              [
                  ["过零率", "极高", "白噪声大幅增加过零次数"],
                  ["频谱质心", "高", "噪声拉高频谱质心"],
                  ["MFCC 统计值", "高", "噪声直接污染 MFCC 系数"],
                  ["带宽", "中高", "噪声展宽频带"],
              ])
    add_para(doc, "L4：特征提取开销反而比 YAMNet 更大", bold=True)
    add_para(doc, "LightGBM 推理延迟 154.5 ms（特征提取 154.3 ms + 推理 0.2 ms），YAMNet 推理延迟 86.0 ms。LightGBM 在速度上也不占优。", indent=True)

    add_heading(doc, "3.6 YAMNet 端到端代码问题（已修复）", 2)
    add_para(doc, "以下问题在 2026-07-16 的代码审查中发现并已修复，详见 YAMNet/src/ 下三个文件。")
    add_table(doc,
              ["编号", "问题", "严重程度", "状态"],
              [
                  ["C1", "类别权重 USE_CLASS_WEIGHTS=True 声明了但从未使用", "P0", "已修复"],
                  ["C2", "MixUp 训练路径无 @tf.function，训练慢 3-10 倍", "P0", "已修复"],
                  ["C3", "选择性解冻失效，整个 YAMNet 被全部解冻", "P0", "已修复"],
                  ["C4", "model.save() 含 hub.KerasLayer，加载需联网可能失败", "P0", "已修复"],
                  ["C5", "噪声评估 RNG 每个 SNR 档重置，与旧版不一致", "P1", "已修复"],
                  ["C6", "学习率衰减只触发一次", "P1", "已修复"],
                  ["C7", "model.predict() 一次性预测全部测试集可能 OOM", "P1", "已修复"],
                  ["C8", "推理测量缺少逐折细节", "P2", "已修复"],
              ])
    add_chart(doc, chart_code_bug_status(), "图 7  YAMNet 端到端代码问题修复状态")

    # Section 4
    add_heading(doc, "四、建议的解决方案", 1)

    add_heading(doc, "4.1 立即可做（不需教授确认）", 2)
    add_table(doc,
              ["编号", "方案", "对应问题", "投入", "预期收益"],
              [
                  ["S1", "YAMNet 端到端实跑", "A4, T1, T2", "Kaggle 2.5-5h GPU", "Clean 准确率 1.91% → 预期 3-8%"],
                  ["S2", "补充 Top-5 / Top-10 准确率", "E1, E3", "本地计算", "Top-5 可能达 10-15%"],
                  ["S3", "按类别频率分组分析", "E2", "本地计算", "Top-50 高频物种可能达 15-25%"],
                  ["S4", "三模型统一对比图", "—", "等 FastAI 结果", "生成合并衰减曲线 + 混淆矩阵"],
              ])

    add_heading(doc, "4.2 建议执行（需投入但效果显著）", 2)
    add_table(doc,
              ["编号", "方案", "对应问题", "投入", "预期收益"],
              [
                  ["S5", "扩大数据集：每类 ≥ 20 样本", "D1, D2, T4", "重新运行数据管道", "准确率可能从约 2% 提升到 10-20%"],
                  ["S6", "Top-50 物种子集评估", "D1, E2", "筛选 + 重跑", "子集上准确率可能达 20-40%"],
                  ["S7", "跑 BirdNET 基线", "—", "Kaggle + birdnetlib", "同台对比，实证优于 BirdNET"],
                  ["S8", "利用 secondary_labels 做多标签训练", "D3", "修改训练代码", "减少标签噪声，提升尾部类表现"],
              ])

    add_heading(doc, "4.3 锦上添花（资源允许时考虑）", 2)
    add_table(doc,
              ["编号", "方案", "对应问题", "投入", "预期收益"],
              [
                  ["S9", "改嵌入聚合为分段预测 + 聚合", "A2", "修改推理代码", "提高有效信号占比"],
                  ["S10", "加深分类头（加 attention 或多层）", "A3", "修改模型结构", "提升分类能力"],
                  ["S11", "峰值归一化改 RMS 归一化", "—", "修改音频预处理", "减少瞬态噪声干扰"],
              ])

    add_heading(doc, "4.4 优先级排序", 2)
    add_para(doc, "必做（项目完整性的前提）：", bold=True)
    for item in [
        "S1：YAMNet 端到端实跑（代码已修复，直接上 Kaggle）",
        "S7：BirdNET 基线对比（「比 BirdNET 好」的实证前提）",
        "S4：三模型统一对比（等 FastAI 结果）",
    ]:
        add_bullet(doc, item)
    add_para(doc, "强烈建议（投入产出比高）：", bold=True)
    for item in [
        "S5 或 S6：扩大数据集 / 子集评估（准确率提升最直接）",
        "S2：Top-5 准确率（零成本，大幅改善结果呈现）",
        "S3：分组分析（零成本，展示部分类别已学到）",
    ]:
        add_bullet(doc, item)
    add_para(doc, "锦上添花：", bold=True)
    for item in [
        "S8：多标签训练",
        "S9-S11：架构优化",
    ]:
        add_bullet(doc, item)

    # Section 5
    add_heading(doc, "五、建议与教授讨论确认的问题", 1)

    discussions = [
        ("问题 1：数据集规模——是否扩大？",
         "背景：当前每类仅约 3.9 个训练样本，是准确率低的根本原因。如果扩到每类 ≥ 20 样本，子集约 24,580 条。",
         ["是否允许扩大抽样子集？扩大到每类多少样本？",
          "如果扩大后计算资源不够，是否可以只选 Top-50 或 Top-100 高频物种？",
          "1229 类全覆盖但每类样本少 vs. 少类但每类样本多，哪个更符合课程要求？"]),
        ("问题 2：评估指标——Top-1 够不够？",
         "背景：1229 类的 Top-1 准确率天然极低，学术界细粒度分类通常用 Top-5。",
         ["报告中是否可以同时报告 Top-1、Top-5、Top-10 准确率？",
          "是否可以加 macro-F1 作为辅助指标？",
          "是否可以按类别频率分组报告？"]),
        ("问题 3：BirdNET 基线——是否必须做？",
         "背景：项目提案中说「比 BirdNET 更好」，但没有同台对比数据。",
         ["BirdNET 基线对比是否为课程硬性要求？",
          "如果时间不够，是否可以用文献中 BirdNET 的公开数据作为引用对比？",
          "如果做，是在全量 1229 类上做，还是只在 Top-50 子集上做？"]),
        ("问题 4：项目范围——1229 类是否过多？",
         "背景：原始提案描述为 12,400 条记录、86 个物种，实际合并后扩大到 1,229 个物种。",
         ["是否可以回到原始提案的约 86 物种方案？",
          "或者筛选 Top-100 / Top-200 高频物种作为主实验？",
          "1229 类的全覆盖是加分项还是负担？"]),
        ("问题 5：端到端微调——预期结果和风险",
         "背景：YAMNet 端到端微调代码已修复，预期 clean 准确率 3-8%，训练时间 30-60 min/fold。",
         ["Kaggle 12 小时 session 限制是否够用？",
          "如果 OOM，batch_size 从 8 降到 4 是否可接受？",
          "如果端到端结果仍然不理想，主要的叙事方向是什么？"]),
        ("问题 6：FastAI 进度——时间线",
         "背景：FastAI CNN 是三模型对比的第三极，目前尚未完成。",
         ["FastAI 部分的预期完成时间？",
          "是否需要协助调试或提供数据接口规范？",
          "如果 FastAI 无法按时完成，两模型对比是否足够？"]),
        ("问题 7：最终报告——叙事方向",
         "背景：在绝对准确率不高（约 2%）的情况下，如何讲好故事是关键。",
         ["报告的核心叙事是「噪声鲁棒性对比」还是「绝对准确率」？",
          "是否可以强调「在标准化噪声条件下系统对比三种异构算法 + BirdNET 基线」作为方法学贡献？",
          "低资源场景的方法论启示是否可以作为结论之一？"]),
    ]
    for title, bg, items in discussions:
        add_para(doc, title, bold=True)
        add_para(doc, bg, indent=True)
        add_para(doc, "需确认：", bold=True)
        for item in items:
            add_bullet(doc, "□ " + item)

    # Section 6
    add_heading(doc, "六、问题优先级总览", 1)
    add_table(doc,
              ["层面", "问题数", "最严重的问题", "可改善性"],
              [
                  ["数据", "5", "D1：每类仅 3-4 样本", "扩大数据集（S5/S6）"],
                  ["模型架构", "4", "A4：冻结编码器无法适配", "端到端微调（S1，已修复待跑）"],
                  ["训练策略", "4", "T1+T2：无增强 + 无类别平衡", "端到端已修复（MixUp + class weight）"],
                  ["评估方法", "3", "E1：Top-1 过于严苛", "补报 Top-5 / 分组分析（S2/S3）"],
                  ["LightGBM 特有", "4", "L1：手工特征无法表达时频结构", "方法论限制，调参改不了"],
                  ["代码 Bug", "8", "C1-C4：影响训练效果和可用性", "全部已修复"],
              ])
    add_chart(doc, chart_problem_overview(), "图 8  各层面发现问题数量")

    add_para(doc, "核心结论：当前结果不理想是「数据极度稀缺 + 模型无法适配 + 无增强无平衡 + 评估方式严苛」四重因素叠加。其中端到端代码问题已修复，数据规模和评估方式需要与教授讨论后决定。", bold=True)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("最后更新：2026年7月16日")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.save(OUTPUT_PATH)
    print(f"已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
